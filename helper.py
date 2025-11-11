import docx
import hashlib
import io
import pdfplumber
import streamlit as st
import tempfile

def file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()

@st.cache_data
def read_docx(file) -> str:
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

@st.cache_data
def read_pdf_chunks(file_bytes: bytes, chunk_size=5000) -> list[str]:
    """
    Return a list of text chunks that are roughly `chunk_size` characters each.
    Keeps page order.
    """
    chunks = []
    current = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            # very long PDFs often contain a lot of whitespace. So collapse it.
            text = " ".join(text.split())
            if len(current) + len(text) > chunk_size:
                if current:
                    chunks.append(current.strip())
                current = text
            else:
                current += " " + text
    
    if current:
        chunks.append(current.strip())

    return chunks

@st.cache_resource
def get_tokenizer():
    """GPT-4o-mini uses the encoding: cl100k_base"""
    import tiktoken
    return tiktoken.encoding_for_model("gpt-4o-mini")

def split_into_token_chunks(text: str, max_tokens: int = 2000) -> list[str]:
    """
    Split a string into chunks that are <= max_tokens tokens long (roughly
    the limit you can send to GPT-4o-mini in a single call). 
    """
    enc = get_tokenizer()
    tokens = enc.encode(text)
    chunks = []
    start = 0
    while start < len(tokens):
        end = min(start + max_tokens, len(tokens))
        chunk_tokens = tokens[start:end]
        chunk_text = enc.decode(chunk_tokens)
        chunks.append(chunk_text)
        start = end
    return chunks

def download_txt(text: str, filename: str = "translated_file.txt"):
    """Creates a download translation as .txt file button."""
    return st.download_button(
        label = "Download as .txt",
        data = text,
        file_name = filename,
        mime="text/plain",
    )

def download_docx(text: str, filename: str = "translated_file.docx"):
    """Creates a download translation as .docx file button."""
    doc = docx.Document()
    for para in text.split("\n"):
        doc.add_paragraph(para)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        with open(tmp.name, "rb") as f:
            data = f.read()
    return st.download_button(
        label = "Download as .docx",
        data = data,
        file_name = filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def perplexity_check(text: str) -> float:
    """Return a crude "perplexity" estimate for the text using a trigram model.
    
    If the perplexity is too high (e.g. >1500), you may flag the translation for user review.
    """
    from collections import defaultdict
    import math
    
    words = text.split()
    trigram_counts = defaultdict(int)
    bigram_counts = defaultdict(int)
    unigram_counts = defaultdict(int)

    for i in range(len(words)):
        unigram_counts[words[i]] += 1
        if i >= 1:
            bigram_counts[(words[i-1], words[i])] += 1
        if i >= 2:
            trigram_counts[(words[i-2], words[i-1], words[i])] += 1

    V = len(unigram_counts)
    N = len(words)

    log_perp = 0.0
    for i in range(2, N):
        trigram = (words[i-2], words[i-1], words[i])
        bigram = (words[i-2], words[i-1])
        count_trigram = trigram_counts[trigram] + 1 # add-one smoothing
        count_bigram = bigram_counts[bigram] + V
        prob = count_trigram / count_bigram
        log_perp += -math.log(prob)

    return math.exp(log_perp / (N - 2))